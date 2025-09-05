import json
from collections import deque
import heapq
import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# ---------- Ticket Class ----------
class Ticket:
    def __init__(self, ticket_id, title, priority="normal", parent=None, status="open"):
        self.ticket_id = ticket_id
        self.title = title
        self.priority = priority
        self.status = status
        self.parent = parent

    def __str__(self):
        return f"[{self.ticket_id}] {self.title} | Status: {self.status} | Priority: {self.priority}"

    def to_dict(self):
        return {
            "ticket_id": self.ticket_id,
            "title": self.title,
            "priority": self.priority,
            "status": self.status,
            "parent": self.parent,
        }

    @staticmethod
    def from_dict(data):
        return Ticket(
            ticket_id=data["ticket_id"],
            title=data["title"],
            priority=data.get("priority", "normal"),
            parent=data.get("parent"),
            status=data.get("status", "open"),
        )

# ---------- Linked List for History ----------
class HistoryNode:
    def __init__(self, action):
        self.action = action
        self.next = None

class TicketHistory:
    def __init__(self):
        self.head = None

    def add_action(self, action):
        new_node = HistoryNode(action)
        if not self.head:
            self.head = new_node
        else:
            current = self.head
            while current.next:
                current = current.next
            current.next = new_node

    def to_list(self):
        actions = []
        current = self.head
        while current:
            actions.append(current.action)
            current = current.next
        return actions

    def load_from_list(self, actions):
        for action in actions:
            self.add_action(action)

    def display_history(self):
        print("\n--- Ticket History (Linked List) ---")
        if not self.head:
            print("No history yet.")
        else:
            current = self.head
            while current:
                print("->", current.action)
                current = current.next
        print("-----------------------------------\n")

# ---------- Help Desk System ----------
class HelpDeskSystem:
    def __init__(self, filename="tickets.json", user=None):
        self.tickets = {}
        self.normal_queue = deque()
        self.priority_queue = []
        self.undo_stack = []
        self.history = TicketHistory()
        self.filename = filename
        self.user = user
        self.load_data()

    # Save tickets + history
    def save_data(self):
        data = {
            "tickets": [t.to_dict() for t in self.tickets.values()],
            "history": self.history.to_list(),
        }
        with open(self.filename, "w") as f:
            json.dump(data, f, indent=4)
        print("Data saved successfully!")

    # Load tickets + history
    def load_data(self):
        try:
            with open(self.filename, "r") as f:
                data = json.load(f)
                for t_data in data.get("tickets", []):
                    ticket = Ticket.from_dict(t_data)
                    self.tickets[ticket.ticket_id] = ticket
                    if ticket.priority == "high" and ticket.status == "open":
                        heapq.heappush(self.priority_queue, (1, ticket.ticket_id))
                    elif ticket.priority == "normal" and ticket.status == "open":
                        self.normal_queue.append(ticket.ticket_id)
                self.history.load_from_list(data.get("history", []))
            print("Data loaded from file.")
        except FileNotFoundError:
            print(" No saved tickets found. Starting fresh.")

    # Analytics
    def analytics_dashboard(self):
        print("\n=== 📊 Analytics Dashboard ===")
        if not self.tickets:
            print("No tickets yet.")
            return
        ticket_matrix = [[t.ticket_id, t.status, t.priority] for t in self.tickets.values()]
        open_count = sum(1 for row in ticket_matrix if row[1] == "open")
        closed_count = sum(1 for row in ticket_matrix if row[1] == "closed")
        high_priority = sum(1 for row in ticket_matrix if row[2] == "high")
        normal_priority = sum(1 for row in ticket_matrix if row[2] == "normal")
        print(f"Total Tickets: {len(ticket_matrix)}")
        print(f"Open Tickets: {open_count} | Closed Tickets: {closed_count}")
        print(f"High Priority: {high_priority} | Normal Priority: {normal_priority}")
        print("==============================\n")

    # Recursion: check dependencies
    def is_resolved(self, ticket_id):
        ticket = self.tickets.get(ticket_id)
        if not ticket:
            return True
        if ticket.status != "closed":
            return False
        if ticket.parent:
            return self.is_resolved(ticket.parent)
        return True

    # CRUD
    def create_ticket(self, ticket_id, title, priority="normal", parent=None):
        if ticket_id in self.tickets:
            print("⚠️ Ticket ID already exists!")
            return
        new_ticket = Ticket(ticket_id, title, priority, parent)
        self.tickets[ticket_id] = new_ticket
        if priority == "high":
            heapq.heappush(self.priority_queue, (1, ticket_id))
        else:
            self.normal_queue.append(ticket_id)
        self.history.add_action(f"Created Ticket {ticket_id}")
        self.undo_stack.append(("delete", ticket_id))
        print(f" Ticket {ticket_id} created successfully!")

    def process_ticket(self):
        if self.priority_queue:
            _, ticket_id = heapq.heappop(self.priority_queue)
        elif self.normal_queue:
            ticket_id = self.normal_queue.popleft()
        else:
            print(" No tickets to process.")
            return
        ticket = self.tickets[ticket_id]
        if not self.is_resolved(ticket.parent):
            print(f"Cannot process Ticket {ticket_id} until parent ticket {ticket.parent} is resolved.")
            return
        ticket.status = "closed"
        self.history.add_action(f"Processed Ticket {ticket_id}")
        self.undo_stack.append(("reopen", ticket_id))
        print(f" Processed Ticket {ticket_id} -> Marked as Closed.")

    def undo(self):
        if not self.undo_stack:
            print(" Nothing to undo.")
            return
        action, ticket_id = self.undo_stack.pop()
        if action == "delete":
            if ticket_id in self.tickets:
                del self.tickets[ticket_id]
            print(f"Undo: Deleted Ticket {ticket_id}")
        elif action == "reopen":
            if ticket_id in self.tickets:
                self.tickets[ticket_id].status = "open"
            print(f" Undo: Reopened Ticket {ticket_id}")
        self.history.add_action(f"Undo {action} on Ticket {ticket_id}")

    def display_tickets(self):
        print("\n--- 🎟️ Current Tickets ---")
        if not self.tickets:
            print("No tickets yet.")
        else:
            for t in self.tickets.values():
                print(t)
        print("---------------------------\n")

    def search_filter(self):
        print("\n=== 🔎 Search & Filter Menu ===")
        print("1. Search by Ticket ID")
        print("2. Filter by Status (open/closed)")
        print("3. Filter by Priority (normal/high)")
        print("4. Back to Main Menu")
        choice = input("Enter choice: ")
        if choice == "1":
            tid = input("Enter Ticket ID: ")
            ticket = self.tickets.get(tid)
            print(ticket if ticket else "⚠️ Ticket not found.")
        elif choice == "2":
            status = input("Enter status (open/closed): ").lower()
            results = [t for t in self.tickets.values() if t.status == status]
            for t in results: print(t) if results else print("⚠️ No tickets found.")
        elif choice == "3":
            priority = input("Enter priority (normal/high): ").lower()
            results = [t for t in self.tickets.values() if t.priority == priority]
            for t in results: print(t) if results else print("⚠️ No tickets found.")
        elif choice == "4":
            return

    # Report Generation
    def generate_report_word(self, filename="HelpDeskReport.docx"):
        doc = Document()
        doc.add_heading("Interactive Help Desk Ticket System Report", 0)
        doc.add_paragraph(f"Generated by: {self.user['username']} ({self.user['role']})")
        doc.add_paragraph(f"Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("------------------------------------------")
        doc.add_heading("Tickets", level=1)
        for t in self.tickets.values():
            doc.add_paragraph(str(t))
        doc.add_heading("Analytics Dashboard", level=1)
        stats = self.analytics_dashboard_stats()
        for k, v in stats.items():
            doc.add_paragraph(f"{k.capitalize()}: {v}")
        doc.add_heading("Ticket History", level=1)
        for action in self.history.to_list():
            doc.add_paragraph("-> " + action)
        doc.save(filename)
        print(f" Word report generated: {filename}")

    def generate_report_pdf(self, filename="HelpDeskReport.pdf"):
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(filename)
        elements = []
        elements.append(Paragraph("Interactive Help Desk Ticket System Report", styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Generated by: {self.user['username']} ({self.user['role']})", styles['Normal']))
        elements.append(Paragraph(f"Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph("Tickets", styles['Heading1']))
        for t in self.tickets.values():
            elements.append(Paragraph(str(t), styles['Normal']))
        elements.append(Paragraph("Analytics Dashboard", styles['Heading1']))
        stats = self.analytics_dashboard_stats()
        for k, v in stats.items():
            elements.append(Paragraph(f"{k.capitalize()}: {v}", styles['Normal']))
        elements.append(Paragraph("Ticket History", styles['Heading1']))
        for action in self.history.to_list():
            elements.append(Paragraph("-> " + action, styles['Normal']))
        doc.build(elements)
        print(f"📄 PDF report generated: {filename}")

    def analytics_dashboard_stats(self):
        ticket_matrix = [[t.ticket_id, t.status, t.priority] for t in self.tickets.values()]
        return {
            "total": len(ticket_matrix),
            "open": sum(1 for row in ticket_matrix if row[1] == "open"),
            "closed": sum(1 for row in ticket_matrix if row[1] == "closed"),
            "high_priority": sum(1 for row in ticket_matrix if row[2] == "high"),
            "normal_priority": sum(1 for row in ticket_matrix if row[2] == "normal"),
        }

# ---------- User Login ----------
def load_users(filename="users.json"):
    try:
        with open(filename, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        default_users = {
            "admin": {"password": "admin123", "role": "admin"},
            "staff": {"password": "staff123", "role": "staff"}
        }
        with open(filename, "w") as f:
            json.dump(default_users, f, indent=4)
        return default_users

def login(users):
    print("\n=== 🔑 Login ===")
    username = input("Username: ")
    password = input("Password: ")
    if username in users and users[username]["password"] == password:
        print(f" Welcome, {username} ({users[username]['role']})!")
        return {"username": username, "role": users[username]["role"]}
    else:
        print(" Invalid credentials.")
        return login(users)

# ---------- Main Menu ----------
def main():
    users = load_users()
    user = login(users)
    system = HelpDeskSystem(user=user)

    while True:
        print("\n=== 🖥️ Help Desk Ticket System ===")
        print("1. Create Ticket")
        print("2. Process Ticket")
        print("3. Undo Last Action")
        print("4. Show All Tickets")
        print("5. Show Analytics Dashboard")
        print("6. Show Ticket History")
        print("7. Save & Exit")
        print("8. Search & Filter Tickets")
        print("9. Generate Report")

        choice = input("Enter choice: ")

        if choice == "1":
            tid = input("Enter Ticket ID: ")
            title = input("Enter Title: ")
            priority = input("Enter Priority (normal/high): ").lower()
            parent = input("Enter Parent Ticket ID (or press Enter if none): ") or None
            system.create_ticket(tid, title, priority, parent)
        elif choice == "2":
            system.process_ticket()
        elif choice == "3":
            system.undo()
        elif choice == "4":
            system.display_tickets()
        elif choice == "5":
            system.analytics_dashboard()
        elif choice == "6":
            system.history.display_history()
        elif choice == "7":
            system.save_data()
            print("👋 Exiting system... Goodbye!")
            break
        elif choice == "8":
            system.search_filter()
        elif choice == "9":
            fmt = input("Generate report as (word/pdf): ").lower()
            if fmt == "word":
                system.generate_report_word()
            elif fmt == "pdf":
                system.generate_report_pdf()
            else:
                print(" Invalid format.")
        else:
            print("Invalid choice. Try again.")

if __name__ == "__main__":
    main()
